from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Add title
title = doc.add_paragraph("Joshua Jensen")
title.runs[0].font.bold = True

# Add contact info
doc.add_paragraph("design@joshjensencreative.com")
doc.add_paragraph("262-271-6729")
doc.add_paragraph("Denver, CO")

# Add About Me
doc.add_paragraph("About Me")
about = doc.add_paragraph("A seasoned Fullstack Developer and UX Designer with a wealth of expertise spanning 14 years, I specialize in crafting exceptional digital experiences. With a strong emphasis on React.js and a keen focus on the cutting-edge Next.js framework, I bring a dynamic skill set to create innovative and user-centric solutions. Beyond my professional pursuits, I thrive on outdoor adventures, including snowboarding, hiking, and exploring off-road terrain with 4x4s and mountain bikes, all while enjoying the company of my spirited husky, Glacier.")

# Add Contact Links
doc.add_paragraph("Contact")
doc.add_paragraph("Website: https://joshjensencreative.com/")
doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/joshua-jensen-30747670/")
doc.add_paragraph("Github: https://github.com/ender2821")

# Add Skills
doc.add_paragraph("Skills")
skills = ["Project Management", "React.js", "Next.js", "Typescript", "Javascript", "Html", "Css", "Sass/Scss", "Tailwind", "Redux", "Apollo Client", "Contentful", "Sanity.io", "GraphQL", ".Net Core", "Sharepoint", "Node.js", "AWS / Vercel", "Azure", "CI/CD with yaml", "LLMs", "MCP Servers", "Git", "MySql", "UX Design", "Graphic Design"]
for skill in skills:
    doc.add_paragraph(skill, style='List Bullet')

# Add Experience
doc.add_paragraph("Experience")

# Frontier Airlines
doc.add_paragraph("Frontier Airlines - Lead Developer")
doc.add_paragraph("Sep 2024 - Present")
exp1 = ["Managed a team of 16 developers based in India and onshore", "Architected and built the new FlyFrontier.com", "Built and integrated Sanity CMS for Enterprise solutions", "Setup CI/CD pipelines for Azure containers and static web apps", "Integrated Generative AI to improve code delivery workflows", "Used LLMs such as Claude with Codepen to efficiently integrate solutions into our code platform", "Integrated several third party SDKs to cut down operating costs, collect analytics, and improve feature delivery", "Added analytics tools such as Application Insights and Noibu to monitor errors and API response times"]
for point in exp1:
    doc.add_paragraph(point, style='List Bullet')

# GAVS Technologies
doc.add_paragraph("GAVS Technologies - Lead Engineer")
doc.add_paragraph("Jan 2024 - Aug 2024")
exp2 = ["Worked as the Lead Web Engineer for Frontier Airlines", "Lead a team of developers based in India and onshore", "Directed knowledge transfer process that included documentation, architecture diagrams, and video presentations", "Built and presented POCs for senior directors from content management systems to caching services", "Integrated Sanity CMS into a Next.js frontend using Typescript and React-query."]
for point in exp2:
    doc.add_paragraph(point, style='List Bullet')

# Portland Webworks
doc.add_paragraph("Portland Webworks - Fullstack Developer Contractor")
doc.add_paragraph("Apr 2023 - Sep 2023")
exp3 = ["Developed front-end components using Next.js", "Constructed mapped data structures integrated into Redux", "Collaborated on API integration through the Java backend", "Styled complex forms utilizing the Tailwind CSS framework"]
for point in exp3:
    doc.add_paragraph(point, style='List Bullet')

# Rightpoint
doc.add_paragraph("Rightpoint - Senior Developer")
doc.add_paragraph("Sep 2018 - Mar 2023")
exp4 = ["Practiced frontend development in HTML, CSS, Sass, JS, and Typescript", "Architected frontends and adhered to UX design standards", "Used React, Next, Node, Sharepoint, Episerver, Git, and Azure", "Led teams and enhanced development practices", "Delivered multiple tech talks, focusing on web application development", "Mentored other developers in UI Development and component library construction with Storybook"]
for point in exp4:
    doc.add_paragraph(point, style='List Bullet')

# Avant
doc.add_paragraph("Avant - UI Developer")
doc.add_paragraph("May 2018 - Aug 2018")
exp5 = ["Wrote HTML and CSS for templates within React", "Conducted cross-browser and WCAG accessibility testing", "Experimented with converting styles into React Styled Components", "Wrote mock data utilizing GraphQL and TypeScript-based schemas"]
for point in exp5:
    doc.add_paragraph(point, style='List Bullet')

# Riskbone
doc.add_paragraph("Riskbone - Product Developer")
doc.add_paragraph("May 2015 - Oct 2017")
exp6 = ["Conceptualized, designed, and developed the Riskbone customer facing website", "Created a trading web app for the Chicago Board of Trade", "Conceptualized, designed, and coded the flagship web app, Level Trading Field", "Managed a team of front-end programmers", "Conducted sales pitches and product demos"]
for point in exp6:
    doc.add_paragraph(point, style='List Bullet')

# Inet
doc.add_paragraph("Inet - Web Design & Developer")
doc.add_paragraph("March 2014 - April 2015")

# Save the document
doc.save('Joshua_Jensen_Resume.docx')
print("Resume created: Joshua_Jensen_Resume.docx")
